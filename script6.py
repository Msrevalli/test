from docx import Document  # Importing the Document class from the docx module
import os  # Importing the os module for interacting with the operating system

# Function to replace a string in a Docx file
def replace_string_in_docx(doc, file_path, old_string, new_string):
    for para in doc.paragraphs:  # Iterate through each paragraph in the document
        if old_string in para.text:  # If the old string is found in the paragraph
            print("Before Updating")
            print(para.text)  # Print the paragraph before updating
            para.text = para.text.replace(para.text, new_string)  # Replace the old string with the new one
            print("After Updating")
            print(para.text)  # Print the paragraph after updating
    doc.save(file_path)  # Save the modified document

# Main function
def main():
    # Ask the user for input directory
    directory = input('Enter Directory: ')

    # Ask the user for the name of the source file and its updated version
    source_file = input("Enter the name of the source file (e.g., Source1.c): ")
    source_new_version = int(input(f"Enter the updated version for {source_file}: "))

    # Ask the user for the name of the test procedure file and its updated version
    test_pr = input("Enter name of test procedure (e.g., TestProcedure1.lua): ")
    test_new_version = int(input(f"Enter the updated version for {test_pr}: "))

    source_file_found = False  # Flag to track if the source file is found in the directory

    # Traverse through the directory and its subdirectories
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):  # If the file is a Docx file
                file_path = os.path.join(root, file)  # Get the full path of the file
                
                doc = Document(file_path)  # Open the Docx file as a Document object
                if source_file in doc.paragraphs[0].text:  # If the source file name is found in the first paragraph
                    source_file_found = True  # Set the flag to indicate source file found
                    # Replace the source file name with its updated version in the Docx file
                    replace_string_in_docx(doc, file_path, source_file, f"{source_file} @ {source_new_version}")
                    # Replace the test procedure name with its updated version in the Docx file
                    replace_string_in_docx(doc, file_path, test_pr, f"{test_pr} @ {test_new_version}")

    if not source_file_found:
        print(f"Source file '{source_file}' not found in the directory.")

# Entry point of the program
if __name__ == "__main__":
    main()  # Call the main function
